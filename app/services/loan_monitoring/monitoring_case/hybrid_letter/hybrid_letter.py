from datetime import datetime, timedelta
from app.models.brief_case.directories.bank_mfo import bank_mfo

from app.models.brief_case.directories.client_region import client_region
from app.models.brief_case.loan_portfolio_schedule import LoanPortfolioSchedule
from app.models.problems_case.problem_state_chain_model import ProblemStateChain
from app.models.problems_case.problems_assets.problems_assets_notification_model import ProblemsStateNotification
from app.services.loan_monitoring.problems_case.state_chains import set_chain_state_for_letter

from .....models.brief_case.directories.local_code import local_code
from .....models.brief_case.loan_portfolio import Loan_Portfolio
from .....models.loan_case.loan_case_model import LoanCase
from .....models.kad_case.kad_case_model import KADCase
from .....models.brief_case.directories.dis_reg_post_codes import post_codes
from .....models.brief_case.directories.client_district import client_district
from .http_client.hybrid_letter_client import HybridLetterClient
from .post_request import HTTPClient
from fastapi import HTTPException
from .....models.monitoring_case.hybrid_letter_model import HybridLetters

from .....common.commit import commit_object, flush_object
from sqlalchemy.orm import aliased
from .....schemas.notification_schemas import CreateNotification
from ...notification.notification_crud import Notificaton_class
from .....common.dictionaries import notification_dictionary
from fastapi.responses import FileResponse

from .....common.dictionaries.monitoring_case_dictionary import letter_status
from .....config.logs_config import info_logger

from .....common.dictionaries.general_tasks_dictionary import JGT, MGT, PGT, GTCAT
from docx import Document
from datetime import date
from sqlalchemy import or_, and_
from io import BytesIO
from barcode import Code128
from barcode.writer import ImageWriter
import subprocess
import base64
def send_letter(request, db_session, loan_id=None):
    exit()
    print("start send")
    print("\n")
    get_letter = db_session.query(HybridLetters).filter(HybridLetters.kad_case_id == request.kad_case_id)\
        .filter(HybridLetters.general_task_id == request.general_task_id).first()
    if get_letter is not None:
        response = send_letter_to_pochtampt(request.loan_portfolio_id, get_letter, request.general_task_id, db_session, loan_id)
    else:
        info_logger.info(f'Letter has already sent for {request.kad_case_id}')
        raise HTTPException(status_code=400, detail='Letter has already sent.')
    
    commit_object(db_session)
    return response

def generate_letter(request, db_session):
    hybrid_letter = None
    file_base_64 = ''
    
    get_hybrid_letter = db_session.query(HybridLetters).filter(HybridLetters.kad_case_id == request.kad_case_id)\
        .filter(HybridLetters.general_task_id == request.general_task_id).first()
        
    if get_hybrid_letter is None:
        new_hybrid_letter = HybridLetters(kad_case_id = request.kad_case_id,
                                            letter_receiver_type_id = request.letter_receiver_type_id,
                                              general_task_id = request.general_task_id,
                                                letter_status_id = letter_status['новый'],
                                                created_at = datetime.now())
        db_session.add(new_hybrid_letter)
        flush_object(db_session)
        
        result = getLetter(request, new_hybrid_letter.general_task.level-1, new_hybrid_letter.id)
        if result == 0:
            info_logger.info(f'borrower type has not found')
            new_hybrid_letter.error_comment = 'Отсутствует тип клиента.'
            new_hybrid_letter.letter_status_id = letter_status['ошибка']
            commit_object(db_session)
        else:
            new_hybrid_letter.letter_base64 = result['base_64']
            new_hybrid_letter.letter_post_id = result['letter_post_id']
            new_hybrid_letter.post_id = result['post_id']
            filename = result['file_path'].replace('project_files/letter_files/','')
            commit_object(db_session)
            file_base_64 = result['base_64']
            hybrid_letter = new_hybrid_letter
            
    else:
        
        result = getLetter(request, get_hybrid_letter.general_task.level-1, get_hybrid_letter.id)
        if result == 0:
            info_logger.info(f'borrower type has not found')
            get_hybrid_letter.error_comment =get_hybrid_letter.error_comment +' ||'+ 'Отсутствует тип клиента.'
            get_hybrid_letter.letter_status_id = letter_status['ошибка']
            commit_object(db_session)
        else:
            get_hybrid_letter.letter_base64 = result['base_64']
            get_hybrid_letter.letter_post_id = result['letter_post_id']
            get_hybrid_letter.post_id = result['post_id']
            get_hybrid_letter.updated_at = datetime.now()
            commit_object(db_session)
            file_base_64 = result['base_64']
            hybrid_letter=get_hybrid_letter
    
    
    
    # bytes = base64.b64decode(result['base_64'], validate=True)

    # f = open('project_files/letter_files/file.pdf', 'wb')
    # f.write(bytes)
    # f.close()
    
    return {"file_base64":file_base_64,
            "letter":hybrid_letter}
    #return FileResponse('project_files/letter_files/file.pdf', filename=filename)


def generate_letter_for_schedule(request, db_session):
    hybrid_letter = None
    file_base_64 = ''
    
    get_hybrid_letter = db_session.query(HybridLetters).filter(HybridLetters.kad_case_id == request.kad_case_id)\
        .filter(HybridLetters.general_task_id == request.general_task_id).first()
        
    if get_hybrid_letter is None:
        new_hybrid_letter = HybridLetters(kad_case_id = request.kad_case_id,
                                            letter_receiver_type_id = request.letter_receiver_type_id,
                                              general_task_id = request.general_task_id,
                                                letter_status_id = letter_status['новый'],
                                                created_at = datetime.now())
        db_session.add(new_hybrid_letter)
        flush_object(db_session)
        
        result = getLetter(request, new_hybrid_letter.general_task.level-1, new_hybrid_letter.id)
        if result == 0:
            info_logger.info(f'borrower type has not found')
            new_hybrid_letter.error_comment = 'Отсутствует тип клиента.'
            new_hybrid_letter.letter_status_id = letter_status['ошибка']
            commit_object(db_session)
        else:
            new_hybrid_letter.letter_base64 = result['base_64']
            new_hybrid_letter.letter_post_id = result['letter_post_id']
            new_hybrid_letter.post_id = result['post_id']
            filename = result['file_path'].replace('project_files/letter_files/','')
            commit_object(db_session)
            file_base_64 = result['base_64']
            hybrid_letter = new_hybrid_letter
            
    elif  get_hybrid_letter.letter_status_id == letter_status['новый']:
        
        result = getLetter(request, get_hybrid_letter.general_task.level-1, get_hybrid_letter.id)
        if result == 0:
            info_logger.info(f'borrower type has not found')
            get_hybrid_letter.error_comment =get_hybrid_letter.error_comment +' ||'+ 'Отсутствует тип клиента.'
            get_hybrid_letter.letter_status_id = letter_status['ошибка']
            commit_object(db_session)
        else:
            get_hybrid_letter.letter_base64 = result['base_64']
            get_hybrid_letter.letter_post_id = result['letter_post_id']
            get_hybrid_letter.post_id = result['post_id']
            get_hybrid_letter.updated_at = datetime.now()
            commit_object(db_session)
            file_base_64 = result['base_64']
            hybrid_letter=get_hybrid_letter
    
    
    
    # bytes = base64.b64decode(result['base_64'], validate=True)

    # f = open('project_files/letter_files/file.pdf', 'wb')
    # f.write(bytes)
    # f.close()
    
    return {"file_base64":file_base_64,
            "letter":hybrid_letter}
    #return FileResponse('project_files/letter_files/file.pdf', filename=filename)






async def append_letter(request, letter, db_session):
    get_letter = db_session.query(HybridLetters).filter(HybridLetters.monitoring_case_id == request.monitoring_case_id)\
        .filter(HybridLetters.general_task_id == request.general_task_id).first()
    
    if get_letter is not None:
        raise HTTPException(status_code=400, detail='Letter has already sent.')
    
    byte_letter = await letter.read()
    base64_letter = base64.b64encode(byte_letter).decode()

    new_hybrid_letter = HybridLetters(  
                                        monitoring_case_id = request.monitoring_case_id,
                                        letter_receiver_type_id = request.letter_receiver_type_id,
                                        general_task_id = request.general_task_id,
                                        post_id = request.post_id,
                                        letter_post_id = "MKB" + request.post_id,
                                        letter_base64 = base64_letter,
                                        send_date = request.send_date,
                                        letter_status_id = letter_status['отправлен'],
                                        created_at = datetime.now(),
                                    )
    
    db_session.add(new_hybrid_letter)

    commit_object(db_session)

    return ["OK"]






def get_post_code_by_loan_portfolio_id(id, db_session):
    
    get_loan_portfolio = db_session.query(Loan_Portfolio.loan_id, Loan_Portfolio.client_name, Loan_Portfolio.client_address, Loan_Portfolio.client_district).filter(Loan_Portfolio.id == id).first()
    
    get_district = db_session.query(client_district).filter(client_district.code == get_loan_portfolio.client_district).first()
    
    if get_district is not None:
        return{'dist_code_post':get_district.post_code and get_district.post_code or None,
           'reg_code_post':get_district.region and get_district.region.post_code and get_district.region.post_code or None,
           'dist_code_iabs':get_district.code,
           'dist_name_iabs':get_district.name,
           'client_name': get_loan_portfolio.client_name,
           'client_address': get_loan_portfolio.client_address,
           'loan_id':get_loan_portfolio.loan_id}
    else: return None
    
    
    
def update_letter_set_repaid_true(letter_id, db_session):
    letter = db_session.query(HybridLetters).filter(HybridLetters.id == letter_id).first()
    letter.is_repaid = True
    flush_object(db_session)
    return letter



def get_hybrid_letter_letter(letter_id, db_session):
    return db_session.query(HybridLetters).filter(HybridLetters.id == letter_id).first()


import uuid



def handleData(letter_number, letter_id):
    # get current date, barcode and letter_post_id and write them to request and response
    cur_date = date.today().strftime("%d.%m.%Y")
    post_id = f'{letter_number}' + (12 - 4 - len(str(letter_id))) * '0' + f'{str(uuid.uuid4().int)[:6]}'
    letter_post_id = 'MKB'  + post_id
    option = {
        "module_width": 0.12,
        "module_height":3.5,
        "font_size": 4,
        "text_distance": 1.8,
        "quiet_zone": 0.7,
    }
    get_barcode = Code128(letter_post_id, writer=ImageWriter())
    
    barcode = BytesIO()
    get_barcode.write(barcode, options=option)
    
    return {"send_date":cur_date,
            "post_id":post_id,
            "letter_post_id":letter_post_id,
            "generate_barcode":barcode}

def readParagraphs(doc_chunk, request, response):
    #find var in doc file and replace them to user data
    for paragraph in doc_chunk:
        hasVar = paragraph.text.find("{$")
        while hasVar != -1:
            left = hasVar + 2
            right = paragraph.text.find("}", hasVar)
            var = paragraph.text[left:right]
            if var == "generate_barcode":
                paragraph.text = ''
                r = paragraph.add_run()
                r.add_picture(response["generate_barcode"])
            else:
                replaceTmp(paragraph, var, request)
            hasVar = paragraph.text.find("{$")
    
    return

def replaceTmp(paragraph, var, request):
    paragraph.text = paragraph.text.replace('{$' + var + '}', str(request[var]))
    return




def generate_pdf(doc_path, path):

    subprocess.call(['soffice','--convert-to','pdf','--outdir',path,doc_path])
    return doc_path


import re
def  getLetter(request, letter_number, letter_id):
    
    borrower_type = None
    docTemplateName = None
    if request.borrower_type is not None:
        borrower_type = request.borrower_type
        is_yurik = re.search('^(?!08|11).*', borrower_type)
        is_fiz = re.search(f'^(08).*', borrower_type)
        is_ip = re.search(f'^(11).*', borrower_type)
        
        if is_yurik:
            docTemplateName = "new_letter_template_business.docx"
        elif is_fiz:
            docTemplateName = "new_letter_template.docx"
        elif is_ip:
            docTemplateName = "new_letter_template_business.docx"
    else:
        return 0
    
    
    pathToTemplateFile = 'project_files/letter_files/'
    docFileName = 'newLetter.docx'
    pdfFileName = 'newLetter.pdf'
    #docTemplateName = "new_letter_template.docx"
    
    doc = Document(pathToTemplateFile + docTemplateName)

    response = handleData(letter_number, letter_id)
    
    request = dict(request)
    
    readParagraphs(doc.paragraphs, request, response)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                readParagraphs(cell.paragraphs,request, response)

    doc.save(pathToTemplateFile + docFileName)
    #got docx file

    in_file = pathToTemplateFile + docFileName
    out_file = pathToTemplateFile + pdfFileName
    
    generate_pdf(in_file, pathToTemplateFile)
    
    with open(out_file, "rb") as pdf_file:
        base64_pdf = base64.b64encode(pdf_file.read()).decode()

    

    return {"send_date":response['send_date'],
            "letter_post_id":response['letter_post_id'],
            "post_id":response['post_id'],
                "base_64": base64_pdf,
                "file_path":out_file
                } 






    

    
    
    
    


def check_expiration(request, db_session):
    get_problems_case = db_session.query(LoanCase).all()
    for problems in get_problems_case:
        if problems.task.general_task_id == 15 or problems.task.general_task_id == 14:
            
            get_letter = db_session.query(HybridLetters)\
                .filter(HybridLetters.monitoring_case_id == request.monitoring_case_id)\
                    .filter(HybridLetters.is_repaid == False)\
                .order_by(HybridLetters.id.desc()).first()
                
            today = datetime.now()    
            n_days_ago = today - timedelta(days=15)
            if n_days_ago > get_letter.send_date:
        
                data = CreateNotification()
                data.from_user_id = problems.task.responsible_id
                data.to_user_id = problems.task.task_manager_id
                data.notification_type = notification_dictionary.notification_type['problem_letter']
                data.body = notification_dictionary.notification_body['send_letter_again']
                data.url = f'{problems.loan_portfolio_id}'+':'+ f'{problems.id}'
                
                notifiaction = Notificaton_class(data)
                notifiaction.create_notification(db_session)
            
            commit_object(db_session)
        
        
        
     
  
def send_letter_to_pochtampt(loan_portfolio_id, get_letter, general_task_id, db_session, loan_id=None):
    print("\n \n")
    info_logger.info(f'started sending letter')
    response = False
    post_codes = get_post_code_by_loan_portfolio_id(loan_portfolio_id, db_session)
    if loan_id is None:
        loan_id= post_codes['loan_id']
    if get_letter.post_id is None or post_codes is None\
        or post_codes['dist_code_post'] is None\
            or post_codes['reg_code_post'] is None\
        or post_codes['client_name'] is None\
            or post_codes['client_address'] is None:
        info_logger.info(f'something went wrong')
        get_letter.error_comment = 'Отсутствуют данные для отправки почты(post code, region code)'
        get_letter.letter_status_id = letter_status['ошибка']
        commit_object(db_session)
        
        #raise HTTPException(status_code=403, detail='Почта не отправлена, данные не хватает!')
    else:
        letter_body = {
                    'Id': get_letter.post_id,
                    'AreaId': int(post_codes['dist_code_post']),
                    'RegionId': int(post_codes['reg_code_post']),
                    'ReceiverFullName': post_codes['client_name'],
                    'ReceiverAddress': post_codes['client_address'],
                    'Base64Content': get_letter.letter_base64
                    }
        
        sending_letter = HybridLetterClient(letter_body)
        request = HTTPClient(sending_letter.send_letter(), sending_letter.get_api_token())
        print('before send')
        response = request.new_post_request()
        print('after send')
        #response = 1
        if not response:
            info_logger.info(f'error while sending letter')
            get_letter.letter_status_id = letter_status['ошибка']
            commit_object(db_session)
            #raise HTTPException(status_code=404, detail='Unable to connect to server api')
        else:
            get_letter.letter_status_id = letter_status['отправлен']
            get_letter.send_date = datetime.now()
            set_chain_state_for_letter(general_task_id, loan_id, db_session)
        get_problems_notification = db_session.query(ProblemsStateNotification).filter(ProblemsStateNotification.loan_portfoilo_id==loan_portfolio_id).first()
        if get_problems_notification is not None:
            if general_task_id == 7:
                get_problems_notification.letter_35_status_id = True
            else:
                get_problems_notification.letter_45_status_id = True
    return response


def send_letter_to_pochtampt_from_non_target(request, get_letter, file_base64, db_session):
    info_logger.info(f'started sending letter')
    response = False
    post_codes = get_post_code_by_loan_portfolio_id(request.loan_portfolio_id, db_session)
    
    if get_letter.post_id is None or post_codes is None\
        or post_codes['dist_code_post'] is None\
            or post_codes['reg_code_post'] is None:
        info_logger.info(f'something went wrong')
        get_letter.error_comment = 'Отсутствуют данные для отправки почты(post code, region code)'
        get_letter.letter_status_id = letter_status['ошибка']
        commit_object(db_session)
        
        #raise HTTPException(status_code=403, detail='Почта не отправлена, данные не хватает!')
    else:
        letter_body = {
                    'Id': get_letter.post_id,
                    'AreaId': int(post_codes['dist_code_post']),
                    'RegionId': int(post_codes['reg_code_post']),
                    'ReceiverFullName': request.client_name,
                    'ReceiverAddress': request.client_address,
                    'Base64Content': file_base64
                    }
        
        
        
        sending_letter = HybridLetterClient(letter_body)
        request = HTTPClient(sending_letter.send_letter(), sending_letter.get_api_token())
        print('before send')
        response = request.new_post_request()
        print('after send')
        #response = 1
        if not response:
            info_logger.info(f'started sending letter')
            get_letter.letter_status_id = letter_status['ошибка']
            commit_object(db_session)
            #raise HTTPException(status_code=404, detail='Unable to connect to server api')
        else:
            get_letter.letter_status_id = letter_status['отправлен']
            get_letter.send_date = datetime.now()
    return response
   
   
   
def get_letter_by_id(id, db_session):
    get_hybrid_letter = db_session.query(HybridLetters).filter(HybridLetters.id == id).first()
    check_letter = HybridLetterClient(letter_id=get_hybrid_letter.post_id)
    request = HTTPClient(check_letter.get_by_id(), check_letter.get_api_token())
    response = request.new_get_request()
    if not response:
        raise HTTPException(status_code=404, detail='Unable to connect to server api')
    get_hybrid_letter.perform = response['Perform']
    get_hybrid_letter.perform_date_time = response['PerformDateTime'] 
    get_hybrid_letter.perform_date_time_str = response['PerformDateTimeStr'] 
    get_hybrid_letter.perform_update_date_time = response['PerformUpdateDateTime'] 
    get_hybrid_letter.perform_update_date_time_str = response['PerformUpdateDateTimeStr'] 
    get_hybrid_letter.note = response['Note'] 
    get_hybrid_letter.courier = response['Courier'] 
    get_hybrid_letter.post_index = response['PostIndex']
    commit_object(db_session)
    
    return {"Perform": get_hybrid_letter.perform,
            "PerformDateTime": get_hybrid_letter.perform_date_time,
            "PerformDateTimeStr": get_hybrid_letter.perform_date_time_str,
            "PerformUpdateDateTime": get_hybrid_letter.perform_update_date_time,
            "perform_update_date_time_str": get_hybrid_letter.perform_update_date_time_str,
            "Note": get_hybrid_letter.note,
            "Courier": get_hybrid_letter.courier,
            "PostIndex": get_hybrid_letter.post_index,
            }
    


def get_base64(letter_id, db_session):
    
    get_base_64 = db_session.query(HybridLetters.letter_base64).filter(HybridLetters.id==letter_id).first()
    if get_base_64 is None:
        raise HTTPException(status_code=404, detail='Документ не найден')
    return get_base_64

 
    
def get_areas():
    kpi_overdue = HybridLetterClient()
    request = HTTPClient(kpi_overdue.get_area(), kpi_overdue.get_api_token())
    response = request.new_get_request()
    if not response:
        raise HTTPException(status_code=404, detail='Unable to connect to server api')
    return response
    
def get_regions():
    kpi_overdue = HybridLetterClient()
    request = HTTPClient(kpi_overdue.get_region(), kpi_overdue.get_api_token())
    response = request.new_get_request()
    if not response:
        raise HTTPException(status_code=404, detail='Unable to connect to server api')
    return response   


def get_letter_by_kad_case_id(kad_case_id, general_task_id, db_session):
    
    hybrid_letter = db_session.query(HybridLetters).filter(HybridLetters.kad_case_id == kad_case_id)\
        .filter(HybridLetters.general_task_id == general_task_id).first()
    letters = {}
    
    if hybrid_letter is not None:
        letters = {"id": hybrid_letter.id,
                    "is_repaid":hybrid_letter.is_repaid,
                    "post_id": hybrid_letter.post_id,
                    "letter_status":hybrid_letter.letter_status_id and hybrid_letter.status,
                    "send_date": hybrid_letter.send_date,
                    "error_comment":hybrid_letter.error_comment,
                    "created_at": hybrid_letter.created_at,
                    "updated_at": hybrid_letter.updated_at,
                    "Perform": hybrid_letter.perform,
                    "PerformDateTime": hybrid_letter.perform_date_time,
                    "PerformDateTimeStr": hybrid_letter.perform_date_time_str,
                    "PerformUpdateDateTime": hybrid_letter.perform_update_date_time,
                    "perform_update_date_time_str": hybrid_letter.perform_update_date_time_str,
                    "Note": hybrid_letter.note,
                    "Courier": hybrid_letter.courier,
                    "PostIndex": hybrid_letter.post_index}
        
    return letters   
    



def get_letter_by_portfolio_id(loan_portfolio_id, general_task_id, db_session):
    hybrid_letter = db_session.query(HybridLetters)\
        .join(KADCase, KADCase.id == HybridLetters.kad_case_id)\
            .join(Loan_Portfolio, Loan_Portfolio.id==KADCase.loan_portfolio_id)\
        .filter(KADCase.loan_portfolio_id == loan_portfolio_id)\
        .filter(HybridLetters.general_task_id == general_task_id).first()
    letters = {}
    
    
    if hybrid_letter is not None:
        letters = {"id": hybrid_letter.id,
                    "is_repaid":hybrid_letter.is_repaid,
                    "post_id": hybrid_letter.post_id,
                    "letter_status":hybrid_letter.letter_status_id and hybrid_letter.status,
                    "send_date": hybrid_letter.send_date,
                    "error_comment":hybrid_letter.error_comment,
                    "created_at": hybrid_letter.created_at,
                    "updated_at": hybrid_letter.updated_at,
                    "Perform": hybrid_letter.perform,
                    "PerformDateTime": hybrid_letter.perform_date_time,
                    "PerformDateTimeStr": hybrid_letter.perform_date_time_str,
                    "PerformUpdateDateTime": hybrid_letter.perform_update_date_time,
                    "perform_update_date_time_str": hybrid_letter.perform_update_date_time_str,
                    "Note": hybrid_letter.note,
                    "Courier": hybrid_letter.courier,
                    "PostIndex": hybrid_letter.post_index}
        
    return letters 








def get_details_for_send(loan_id, db_session):
    get_loan = db_session.query(local_code.name.label('local_name'),
                                Loan_Portfolio.client_name, 
                                client_region.name.label('region'),
                                client_district.name.label('client_district'),
                                Loan_Portfolio.client_address,
                                Loan_Portfolio.issue_date,
                                Loan_Portfolio.maturity_date,
                                Loan_Portfolio.loan_id,
                                Loan_Portfolio.credit_line_purpose.label('loan_purpose'),
                                Loan_Portfolio.osn_cmp_percent.label('percent'),
                                Loan_Portfolio.contract_amount_uz_currency.label('loan_amount'),
                                Loan_Portfolio.total_overdue,
                                Loan_Portfolio.overdue_balance,
                                Loan_Portfolio.balance_16377
                                )\
                                .join(local_code, local_code.id == Loan_Portfolio.local_code_id, isouter=True)\
                                .join(client_region, client_region.id == Loan_Portfolio.client_region, isouter=True)\
                                .join(client_district, client_district.code == Loan_Portfolio.client_district, isouter=True)\
                                .filter(Loan_Portfolio.loan_id == loan_id).first()\
                                    
    
    get_schedule = db_session.query(LoanPortfolioSchedule.date_red).filter(LoanPortfolioSchedule.loan_id == loan_id).filter(LoanPortfolioSchedule.date_red != None).first()
    schedule_day = None
    if get_schedule is not None:
        schedule_day = get_schedule.date_red.day
    letter_details = {}
    issue_date = None
    maturity_date = None
    loan_period = None
    
    
    if get_loan is not None:
        if get_loan.issue_date is not None:
            issue_date = get_loan.issue_date
        if get_loan.maturity_date is not None:
            maturity_date = get_loan.maturity_date
        if issue_date is not None and maturity_date is not None:
            loan_period = diff_month(maturity_date, issue_date)
            
        letter_details = {  "local_name":get_loan.local_name,
                            "client_name":get_loan.client_name, 
                            "client_region":get_loan.region,
                            "branch_region":get_loan.region,  
                            "client_district":get_loan.client_district,
                            "client_address":get_loan.client_address,
                            "issue_date":get_loan.issue_date,
                            "loan_period": loan_period,
                            "schedule_day":schedule_day,
                            "loan_id":get_loan.loan_id,
                            "loan_purpose":get_loan.loan_purpose,
                            "percent":get_loan.percent,
                            "loan_amount":get_loan.loan_amount,
                            "total_overdue":get_loan.total_overdue,
                            "overdue_balance":get_loan.overdue_balance,
                            "balance_16377":get_loan.balance_16377,
        }
        
    return letter_details  


def diff_month(d1, d2):
        return (d1.year - d2.year) * 12 + d1.month - d2.month
  
    
    
def get_letters_by_monitoring_id(monitoring_case_id, db_session):
    problems_letter = db_session.query(HybridLetters).filter(HybridLetters.monitoring_case_id == monitoring_case_id).all()
    letters = []
    for letter in problems_letter:
        letters.append({"id": letter.id and letter.id or None,
                    "letter_post_id": letter.letter_post_id and letter.letter_post_id or None,
                    "letter_status":letter.letter_status_id and letter.status or None,
                    "send_date": letter.send_date and letter.send_date or None,
                    "created_at": letter.created_at and letter.created_at or None,
                    "updated_at": letter.updated_at and letter.updated_at or None})
        
    return letters   





def match_post_codes_to_region(db_session):
    
    get_client_district = db_session.query(client_district).all()
    
    for district in get_client_district:
        
        get_post_code = db_session.query(post_codes).filter(district.post_code == post_codes.dist_code_post).first()
        
        if get_post_code is not None:
            get_region = db_session.query(client_region).filter(client_region.post_code == get_post_code.reg_code_post).first()
            
            if get_region is not None:
                district.region_id=get_region.id
                
                flush_object(db_session)
    
    
    commit_object(db_session)
    
    return "OK"






# def set_chain_state_for_letter(general_task_id, loan_id, db_session):
#     set_chain = db_session.query(ProblemStateChain).filter(ProblemStateChain.loan_id== loan_id).first()
    
    
#     if general_task_id==7:
#         state_type=1
#         if set_chain is None:
        
#             new_chain = ProblemStateChain(loan_id=loan_id,
#                                         letter_35_id = state_type,
#                                         letter_35_date =datetime.now(),
#                                         last_state_id = state_type,
#                                         created_at = datetime.now())
#             db_session.add(new_chain)
#             flush_object(db_session)
    
#         else:
#             set_chain.letter_35_id = state_type
#             set_chain.letter_35_date = datetime.now()
#             set_chain.last_state_id = state_type
#             set_chain.updated_at = datetime.now()
            
#             flush_object(db_session)
#     else:
#         state_type=2
#         if set_chain is None:
        
#             new_chain = ProblemStateChain(loan_id=loan_id,
#                                         letter_45_id = state_type,
#                                         letter_45_date =datetime.now(),
#                                         last_state_id = state_type,
#                                         created_at = datetime.now())
#             db_session.add(new_chain)
#             flush_object(db_session)
    
#         else:
#             set_chain.letter_45_id = state_type
#             set_chain.letter_45_date = datetime.now()
#             set_chain.last_state_id = state_type
#             set_chain.updated_at = datetime.now()
            
#             flush_object(db_session)
    
    