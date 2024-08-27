
import datetime
from docx import Document
from docx.shared import Pt
import math
from datetime import date
from ....models.brief_case.loan_portfolio import Loan_Portfolio
from .script_num_to_word import num2word
import subprocess
import base64





def getTemplate(request, db_session):
    
    pathToTemplateFile = 'project_files/monitoring_templates/'
    docFileName = 'newTarget.docx'
    pdfFileName = 'newTarget.pdf'
    docTemplateName = "target_template.docx"
    get_portfolio = db_session.query(Loan_Portfolio).filter(Loan_Portfolio.id == request.loan_portfolio_id).first()
    #get_portfolio = db_session.query(Loan_Portfolio).first()
    
    doc = Document(pathToTemplateFile + docTemplateName)

    data = {}
    #print(round(float(get_portfolio.balance_16379)/1000000, 2))
    data = {
            "loanId": get_portfolio.loan_id,
            "branchDistrict": get_portfolio.client_district,
            "mfo": get_portfolio.local_code.code,
            "currentDate": datetime.datetime.now().date(),
            "branchRegion": get_portfolio.region.name,
            "clientName": get_portfolio.client_name,
            "issueDate": get_portfolio.issue_date,
            "loanTerm": get_portfolio.maturity_date,
            "loanPercent": get_portfolio.osn_cmp_percent and get_portfolio.osn_cmp_percent or 0,
            "creditAccBalance": get_portfolio.credit_account_balance and round(float(get_portfolio.credit_account_balance)/1000000, 2) or 0,
            "clientAddress": get_portfolio.client_address and get_portfolio.client_address or 0,
            "clientInn": get_portfolio.inn_passport and get_portfolio.inn_passport or 0,
            "totalProblemDebt": get_portfolio.total_overdue and round(float(get_portfolio.total_overdue)/1000000, 2) or 0,
            "overdueBalance": get_portfolio.overdue_balance and round(float(get_portfolio.overdue_balance)/1000000, 2) or 0,
            "balance16377": get_portfolio.balance_16377 and round(float(get_portfolio.balance_16377)/1000000, 2) or 0,
            "judicialBalance": get_portfolio.judicial_balance and round(float(get_portfolio.judicial_balance)/1000000, 2) or 0,
            "balance91501": get_portfolio.balance_91501 and round(float(get_portfolio.balance_91501)/1000000, 2) or 0,
            "balance91503": get_portfolio.balance_91503 and round(float(get_portfolio.balance_91503)/1000000, 2) or 0,
            "balance95413": get_portfolio.balance_95413 and round(float(get_portfolio.balance_95413)/1000000, 2) or 0,
            "balance16397": get_portfolio.balance_16397 and round(float(get_portfolio.balance_16397)/1000000, 2) or 0,
            "balance16379": get_portfolio.balance_16379 and round(float(get_portfolio.balance_16379)/1000000, 2) or 0,
            "overdueDays": datetime.datetime.now().date(),
            "overduePercentDays": datetime.datetime.now().date(),
            "loanAccount": get_portfolio.loan_account,
            "loanProduct": get_portfolio.loan_product,
            "creditLinePurpose": get_portfolio.credit_line_purpose,
            "loanAmount": get_portfolio.contract_amount_uz_currency and f'{math.ceil(float(get_portfolio.contract_amount_uz_currency))}  ({num2word(math.ceil(float(get_portfolio.contract_amount_uz_currency)))})' ,
            "currency": get_portfolio.currency.name,
            "maturityDate": get_portfolio.maturity_date,
            "securityDescription": get_portfolio.security_description,
            "clientName": get_portfolio.client_name
            }
    
    
    readParagraphs(doc.paragraphs,data)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                readParagraphs(cell.paragraphs,data)

    doc.save(pathToTemplateFile + docFileName)
    #got docx file

    in_file = pathToTemplateFile + docFileName
    out_file = pathToTemplateFile + pdfFileName
    
    generate_pdf(in_file, pathToTemplateFile)
    
    with open(out_file, "rb") as pdf_file:
        base64_pdf = base64.b64encode(pdf_file.read()).decode()
    #got base64

    

    return {"file_url": in_file
                } 
    
    
    
    
def readParagraphs(doc_chunk,request):
    #find var in doc file and replace them to user data
    for paragraph in doc_chunk:
        hasVar = paragraph.text.find("{$")
        while hasVar != -1:
            left = hasVar + 2 
            right = paragraph.text.find("}", hasVar)
            var = paragraph.text[left:right]
            if var == "generate_barcode":
                paragraph.text = ''
                r = paragraph.add_run()
                #r.add_picture(response["generate_barcode"])
            else:
                replaceTmp(paragraph, var, request)
            hasVar = paragraph.text.find("{$")
    
    return

def replaceTmp(paragraph, var, request):
    paragraph.text = paragraph.text.replace('{$' + var + '}', str(request[var]))
    
    return




def generate_pdf(doc_path, path):

    subprocess.call(['soffice',
                 # '--headless',
                 '--convert-to',
                 'pdf',
                 '--outdir',
                 path,
                 doc_path])
    return doc_path
